#!/usr/bin/env python3
"""
Google Classroom Buddy - Merged Tool
==================================================

This single CLI tool connects to the Google Classroom and Drive APIs to:
1.  (detect-materials) Detect new lecture materials (PDFs, Docs), generate
    study aids (Audio, Flashcards, Quizzes), and upload them to Google Drive.
2.  (detect-announcements) Detect new announcements, scan for keywords
    (project, lab test), and generate tailored project ideas or practice questions.
3.  (summarize-announcements) Provide a single AI summary for all recent
    announcements in a course.
4.  (analyze-announcement) Manually analyze any piece of text for
    project/lab test content.
5.  (list-courses) List all your active courses.
6.  (generate-doc) Scan a source code directory and generate a formatted
    .docx document from the code files.

Setup:
-------
1. Enable Google Classroom & Drive APIs in Google Cloud Console.
2. Download OAuth credentials.json.
3. Install dependencies:
   pip install google-auth google-auth-oauthlib google-api-python-client \
               google-generativeai python-dotenv typer rich pdfplumber gTTS \
               python-docx rich.markdown
4. Create a .env file with:
   GEMINI_API_KEY=your_api_key_here
5. Run for the first time to authenticate (this will request all needed permissions):
   python merged_buddy.py list-courses
6. Run other commands:
   python merged_buddy.py detect-materials --all-courses --since 24
   python merged_buddy.py detect-announcements --all-courses --since 72
   python merged_buddy.py generate-doc --source ./my-project --title "My Project Documentation"
"""

import os
import pickle
import re
import io
from datetime import datetime, timedelta, timezone
from typing import List, Dict, Tuple, Optional
import typer
from rich import print
from rich.console import Console
from rich.panel import Panel
from rich.status import Status
from rich.markdown import Markdown

# Google API Imports
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from googleapiclient.http import MediaIoBaseDownload, MediaIoBaseUpload

# Text & AI Imports
import google.generativeai as genai
from dotenv import load_dotenv
import pdfplumber
from gtts import gTTS

# --- NEW: Import for Docx Generation ---
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
# ------------------------------------

# ---------------------- Setup ----------------------
app = typer.Typer(help="A Google Classroom CLI tool to detect materials, analyze announcements, and generate study aids with Gemini AI.")
console = Console()
load_dotenv()

# --- MERGED SCOPES ---
# We need all permissions from both scripts
SCOPES = [
    'https://www.googleapis.com/auth/classroom.courses.readonly',
    'https://www.googleapis.com/auth/classroom.courseworkmaterials.readonly', # For Study Buddy
    'https://www.googleapis.com/auth/classroom.announcements.readonly',     # For Monitor
    'https://www.googleapis.com/auth/drive.readonly',                       # For Study Buddy
    'https://www.googleapis.com/auth/drive.file'                            # For Study Buddy
]

# --- Keywords from Monitor Script ---
PROJECT_KEYWORDS = [
    'synopsis', 'project', 'pbl', 'problem-based learning', 'problem based learning',
    'capstone', 'thesis', 'research', 'presentation', 'proposal', 'development', 'design',
    'prototype', 'deliverable',
    'milestone', 'case study', 'report', 'documentation', 'mini project',
    'major project', 'final year project', 'term project', 'semester project'
]

LAB_TEST_KEYWORDS = [
    'lab test', 'evaluation', 'practical exam', 'viva', 'quiz',
    'midterm', 'final exam', 'coding test', 'assessment', 'lab evaluation'
]

# ---------------------- Core Class (Merged) ----------------------
class ClassroomBuddyCLI:
    def __init__(self, credentials_file='credentials.json', token_file='token.pickle'):
        self.credentials_file = credentials_file
        self.token_file = token_file
        self.creds = None
        self.classroom_service = None # For Classroom API
        self.drive_service = None     # For Drive API
        self.gemini_model = None
        self._authenticate()
        self._setup_gemini()

    def _authenticate(self):
        """Authenticate with Google APIs (Classroom & Drive)."""
        creds = None
        if os.path.exists(self.token_file):
            with open(self.token_file, 'rb') as token:
                creds = pickle.load(token)

        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                console.print("[yellow]Refreshing expired credentials...[/yellow]")
                try:
                    creds.refresh(Request())
                except Exception as e:
                    console.print(f"[red]Error refreshing token: {e}[/red]")
                    console.print("[yellow]Please re-authenticate.[/yellow]")
                    creds = None # Force re-authentication
                    if os.path.exists(self.token_file):
                         os.remove(self.token_file) # Delete bad token
            
            if not creds:
                if not os.path.exists(self.credentials_file):
                    raise FileNotFoundError(
                        f"Credentials file '{self.credentials_file}' not found. "
                        "Download it from Google Cloud Console."
                    )
                console.print("[cyan]Authenticating with Google...[/cyan]")
                flow = InstalledAppFlow.from_client_secrets_file(self.credentials_file, SCOPES)
                creds = flow.run_local_server(port=0)

            with open(self.token_file, 'wb') as token:
                pickle.dump(creds, token)
            console.print("[green]✓ Authentication successful![/green]")

        self.creds = creds
        try:
            # Build BOTH services
            self.classroom_service = build('classroom', 'v1', credentials=creds)
            self.drive_service = build('drive', 'v3', credentials=creds)
            console.print("[green]✓ Classroom and Drive services initialized.[/green]")
        except HttpError as e:
            console.print(f"[red]Error building services: {e}[/red]")
            raise

    def _setup_gemini(self):
        """Setup Gemini AI."""
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError("Missing GEMINI_API_KEY in environment variables (.env file).")
        genai.configure(api_key=api_key)
        self.gemini_model = genai.GenerativeModel('gemini-2.5-flash')
        console.print("[green]✓ Gemini AI configured![/green]")

    def _parse_timestamp(self, ts: str):
        """Parses Google API timestamp to a comparable datetime object."""
        if not ts:
            return datetime.min.replace(tzinfo=timezone.utc)
        # Handle nanoseconds and 'Z'
        if '.' in ts:
            ts = ts.split('.')[0] + 'Z'
        return datetime.strptime(ts, '%Y-%m-%dT%H:%M:%SZ').replace(tzinfo=timezone.utc)

    # --- Course & Data Fetching Methods ---

    def get_courses(self):
        """Fetch available Google Classroom courses."""
        try:
            results = self.classroom_service.courses().list(
                pageSize=100, courseStates=['ACTIVE']
            ).execute()
            return results.get('courses', [])
        except HttpError as e:
            console.print(f"[red]⚠️ Error fetching courses:[/red] {e}")
            return []

    def get_new_materials(self, course_id, since_hours):
        """
        Fetch new courseWorkMaterial items (lectures, readings, etc.)
        """
        try:
            cutoff_time = datetime.now(timezone.utc) - timedelta(hours=since_hours)
            
            materials = []
            page_token = None
            while True:
                response = self.classroom_service.courses().courseWorkMaterials().list(
                    courseId=course_id,
                    pageSize=20,
                    pageToken=page_token,
                    orderBy='updateTime desc'
                ).execute()
                
                items = response.get('courseWorkMaterial', [])
                
                for item in items:
                    update_time = self._parse_timestamp(item.get('updateTime'))
                    if update_time > cutoff_time:
                        materials.append(item)
                    else:
                        # Stop paging once we hit old items
                        return materials
                
                page_token = response.get('nextPageToken')
                if not page_token or not items:
                    break
                    
            return materials
        except HttpError as e:
            console.print(f"[red]⚠️ Error fetching materials for course {course_id}:[/red] {e}")
            return []

    def get_announcements(self, course_id, max_results=10, since_hours=None):
        """Retrieve course announcements."""
        try:
            announcements, page_token = [], None
            cutoff_time = None
            if since_hours:
                # Use timezone-aware cutoff
                cutoff_time = datetime.now(timezone.utc) - timedelta(hours=since_hours)
                
            while True:
                response = self.classroom_service.courses().announcements().list(
                    courseId=course_id,
                    pageSize=min(max_results, 100),
                    pageToken=page_token,
                    orderBy='updateTime desc'
                ).execute()

                items = response.get('announcements', [])
                
                if not items:
                    break

                if since_hours:
                    filtered_items = []
                    for i in items:
                        # Use the robust, timezone-aware parser
                        update_time = self._parse_timestamp(i.get('updateTime'))
                        if update_time > cutoff_time:
                            filtered_items.append(i)
                        else:
                            # Since announcements are sorted by updateTime desc,
                            # we can stop once we hit an old one.
                            items = filtered_items # Use only the items we found
                            page_token = None # Force stop paging
                            break
                    items = filtered_items
                
                announcements.extend(items)
                page_token = response.get('nextPageToken')
                
                if not page_token or len(announcements) >= max_results:
                    break

            return announcements[:max_results]
        except HttpError as e:
            console.print(f"[red]⚠️ Error fetching announcements:[/red] {e}")
            return []

    # --- Drive & File Handling Methods (from Study Buddy) ---

    def get_drive_file_text(self, drive_file) -> Optional[str]:
        """
        Downloads a Google Drive file (Doc, PDF) and extracts all text.
        """
        file_id = drive_file.get('id')
        name = drive_file.get('title', 'Unknown File')
        
        # Get file metadata to check its MIME type
        try:
            metadata = self.drive_service.files().get(fileId=file_id, fields='mimeType, name').execute()
            mime_type = metadata.get('mimeType')
            name = metadata.get('name', name)
            
            console.print(f"  > Reading file: [dim]{name}[/dim] (Type: {mime_type})")
            
            request = None
            if 'google-apps.document' in mime_type:
                # It's a Google Doc, export as plain text
                request = self.drive_service.files().export_media(
                    fileId=file_id, mimeType='text/plain'
                )
            elif 'google-apps.presentation' in mime_type:
                # It's a Google Slide, export as plain text
                request = self.drive_service.files().export_media(
                    fileId=file_id, mimeType='text/plain'
                )
            elif 'pdf' in mime_type:
                # It's a PDF, download directly
                request = self.drive_service.files().get_media(fileId=file_id)
            else:
                console.print(f"  [yellow]Skipping unsupported file type: {mime_type}[/yellow]")
                return None

            # Download the file content
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            
            fh.seek(0)
            
            # Extract text based on original type
            if 'pdf' in mime_type:
                try:
                    with pdfplumber.open(fh) as pdf:
                        full_text = "\n".join(
                            page.extract_text() for page in pdf.pages if page.extract_text()
                        )
                        return full_text
                except Exception as e:
                    console.print(f"  [red]Error reading PDF content: {e}[/red]")
                    return None
            else:
                # It was exported as plain text
                return fh.read().decode('utf-8')

        except HttpError as e:
            console.print(f"  [red]Error accessing Drive file {name}: {e}[/red]")
            return None

    def _upload_to_drive(self, content: str, filename: str, mime_type: str = 'text/markdown') -> Optional[str]:
        """
        Uploads the generated text content to Google Drive.
        """
        file_metadata = {'name': filename}
        media = MediaIoBaseUpload(
            io.BytesIO(content.encode('utf-8')),
            mimetype=mime_type,
            resumable=True
        )
        try:
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
            return file.get('webViewLink')
        except HttpError as e:
            console.print(f"[red]Error uploading {filename} to Drive: {e}[/red]")
            return None

    def _upload_audio_to_drive(self, audio_bytes_io: io.BytesIO, filename: str) -> Optional[str]:
        """
        Uploads generated audio bytes to Google Drive.
        """
        file_metadata = {'name': filename}
        audio_bytes_io.seek(0) # Rewind the in-memory file
        media = MediaIoBaseUpload(
            audio_bytes_io,
            mimetype='audio/mpeg',
            resumable=True
        )
        try:
            file = self.drive_service.files().create(
                body=file_metadata,
                media_body=media,
                fields='id, webViewLink'
            ).execute()
            return file.get('webViewLink')
        except HttpError as e:
            console.print(f"[red]Error uploading {filename} to Drive: {e}[/red]")
            return None

    # --- Gemini Generation: Study Aids (from Study Buddy) ---

    def _run_gemini_prompt(self, prompt: str, lecture_text: str) -> str:
        """Helper function to run a generation prompt with error handling."""
        full_prompt = f"{prompt}\n\nLECTURE TEXT:\n---\n{lecture_text}\n---"
        try:
            response = self.gemini_model.generate_content(full_prompt)
            return response.text.strip()
        except Exception as e:
            console.print(f"[red]Error during Gemini generation: {e}[/red]")
            return "Error: Could not generate content."

    def generate_audio_narration(self, text: str) -> str:
        prompt = """
        You are a narrator. Based on the provided lecture text, write a
        concise 2-3 minute audio script (narration only).
        This script will be used for a text-to-speech audio summary.
        Focus on the main topics, key definitions, and conclusions.
        Be clear, concise, and speak directly to the listener.
        
        IMPORTANT: Do NOT include any visual cues, headings, titles, or Markdown formatting.
        Output only the plain text of the narration.
        
        Example:
        Hello, and welcome to this lecture summary. Today we're covering three main points. First, what is...
        """
        return self._run_gemini_prompt(prompt, text)

    def generate_flashcards(self, text: str) -> str:
        prompt = """
        You are a study aid generator. Based on the provided lecture text,
        generate 15-20 flashcards in CSV format.
        Each row should be a "Question,Answer" pair.
        Do NOT include a header row.
        Ensure questions are clear and answers are concise.
        
        Example:
        "What is the capital of France?","Paris"
        "What is 2+2?","4"
        """
        return self._run_gemini_prompt(prompt, text)

    def generate_quiz(self, text: str) -> str:
        prompt = """
        You are a professor. Based on the provided lecture text, create a
        10-question multiple-choice quiz in Markdown format.
        For each question, provide 4 options (A, B, C, D).
        At the end of each question, clearly indicate the correct answer.
        
        Example:
        **1. What is the capital of France?**
        A) London
        B) Berlin
        C) Paris
        D) Madrid
        
        *Correct Answer: C*
        """
        return self._run_gemini_prompt(prompt, text)
        
    # --- Gemini Generation: Announcement Analysis (from Monitor) ---

    def summarize_course_announcements(self, course_name, announcements):
        """Summarize all announcements in a single Gemini summary."""
        if not announcements:
            return "No recent announcements to summarize."

        compiled_texts = []
        for ann in announcements:
            # Use the robust parser
            timestamp = self._parse_timestamp(ann.get('updateTime', '')).strftime("%Y-%m-%d %H:%M")
            text = ann.get('text', '').strip() or 'No content.'
            compiled_texts.append(f"[{timestamp}] {text}")

        prompt = f"""
You are a helpful assistant summarizing all recent announcements for a course.

Course: {course_name}

Announcements:
{chr(10).join(compiled_texts)}

Provide a **very short, high-level summary** in bullet points (3-5 points maximum).
Focus *only* on the most critical information:
- **Key deadlines**
- **Required actions** (e.g., "Submit X," "Prepare for Y")
- **Main topics** (e.g., "Project synopsis guidelines released")

Be as brief as possible. Do not use bold text.
"""
        try:
            response = self.gemini_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return f"⚠️ Error generating summary: {str(e)}"

    def detect_project_announcements(self, announcements: List[Dict]) -> List[Tuple[Dict, List[str]]]:
        """
        Detect announcements containing project-related keywords.
        Returns list of tuples: (announcement, matched_keywords)
        """
        project_announcements = []
        
        for ann in announcements:
            text = ann.get('text', '').lower()
            matched_keywords = []
            
            for keyword in PROJECT_KEYWORDS:
                if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
                    matched_keywords.append(keyword)
            
            if matched_keywords:
                project_announcements.append((ann, list(set(matched_keywords))))
        
        return project_announcements

    def detect_lab_test_announcements(self, announcements: List[Dict]) -> List[Tuple[Dict, List[str]]]:
        """
        Detect announcements containing lab test/evaluation keywords.
        Returns list of tuples: (announcement, matched_keywords)
        """
        lab_test_announcements = []
        
        for ann in announcements:
            text = ann.get('text', '').lower()
            matched_keywords = []
            
            for keyword in LAB_TEST_KEYWORDS:
                if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
                    matched_keywords.append(keyword)
            
            if matched_keywords:
                # To avoid overlap, check if it's *also* a project announcement.
                # If it has project keywords, we'll let the project detector handle it.
                is_project = False
                for proj_keyword in PROJECT_KEYWORDS:
                        if re.search(r'\b' + re.escape(proj_keyword) + r'\b', text, re.IGNORECASE):
                            is_project = True
                            break
                
                # Only add if it's NOT primarily a project announcement
                if not is_project:
                    lab_test_announcements.append((ann, list(set(matched_keywords))))
        
        return lab_test_announcements

    def generate_tailored_project_ideas(self, course_name: str, announcement_text: str, keywords: List[str]) -> str:
        """
        Generate project ideas tailored to the specific announcement requirements.
        """
        prompt = f"""
You are an expert educational advisor analyzing a project announcement and generating tailored project ideas.

Course: {course_name}
Project Announcement (Full Text):
{announcement_text}

Detected Keywords: {', '.join(keywords)}

Your task:
Be detailed but concise. Use bullet points heavily.

1. CAREFULLY READ and ANALYZE the announcement to understand:
   - Specific requirements and constraints
   - Topics or domains mentioned
   - Technologies or tools specified
   - Learning objectives
   - Deliverables expected
   - Any deadlines or milestones
   - Team size or collaboration requirements
   - Evaluation criteria

2. Generate 5-7 TAILORED PROJECT IDEAS that:
   - DIRECTLY match the announcement specifications
   - Address the stated requirements
   - Are feasible within the given constraints
   - Align with the course subject and level
   - Incorporate mentioned technologies/tools

3. For EACH project idea provide:
   - **Project Title**: Clear, descriptive name
   - **Description**: 2-3 sentences explaining the project
   - **How it meets requirements**: Explicitly state which announcement requirements it fulfills
   - **Key Technologies/Tools**: Specific tech stack
   - **Implementation Steps**: 4-6 high-level steps
   - **Expected Outcomes**: What students will learn/achieve
   - **Complexity Level**: Beginner/Intermediate/Advanced

4. RESOURCES & REFERENCES:
   - Provide 5-8 specific resources for EACH major technology/tool mentioned
   - Include actual GitHub repository search terms (e.g., "search GitHub for: 'student management system python django'")
   - List relevant tutorial websites with search queries
   - Suggest YouTube channels or specific video searches
   - Recommend documentation links (provide exact URLs where possible)
   - Mention similar projects on platforms like:
     * GitHub (provide search terms)
     * Kaggle (for data science projects)
     * CodePen/JSFiddle (for web projects)
     * Instructables (for hardware projects)

5. EXAMPLE PROJECT LINKS & SEARCH STRATEGIES:
   - Provide specific GitHub search queries that will find similar completed projects
   - Format: "GitHub Search: [exact search term]"
   - Example: "GitHub Search: 'e-commerce website react node mongodb'"
   - Include alternative search terms for different platforms

Format your response clearly with:
- Main headers (##)
- Subheaders (###)
- Bullet points for lists
- Code blocks for search terms or commands
- Bold for emphasis

Be SPECIFIC, PRACTICAL, and ensure all ideas are directly relevant to the announcement content.
"""
        
        try:
            response = self.gemini_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return f"⚠️ Error generating project ideas: {str(e)}"

    def generate_practice_questions(self, course_name: str, announcement_text: str, keywords: List[str]) -> str:
        """
        Generate practice lab test questions based on the announcement.
        """
        prompt = f"""
You are an expert Computer Science professor creating a practice test.

Course: {course_name}
Announcement (Full Text):
{announcement_text}

Detected Keywords: {', '.join(keywords)}

Your task:
Be concise and clear.

1. CAREFULLY READ and ANALYZE the announcement to identify:
   - The key topics or concepts to be tested (e.g., "Data Structures," "Algorithms," "Database Queries," "Python Basics").
   - The format of the test (e.g., coding questions, viva, multiple choice).
   - Any specific technologies or languages mentioned.

2. Generate 5-7 TAILORED PRACTICE QUESTIONS that:
   - DIRECTLY relate to the topics in the announcement.
   - Are at an appropriate difficulty level for a university course.
   - Mimic the likely format of the test (focus on coding problems if it's a "lab test").

3. For EACH practice question, provide:
   - **Question Title/Topic**: (e.g., "Array Manipulation," "SQL Join," "Binary Tree Traversal")
   - **Problem Statement**: A clear, concise problem.
   - **Example Input/Output**: (if applicable)
   - **Key Concepts to Apply**: What the student needs to know to solve it.
   - **Hint**: (Optional) A small hint to guide the student.

4. Provide a "Study Guide & Resources" section:
   - List the 3-5 most important topics to review.
   - Provide 5-8 specific resources (documentation links, tutorials, YouTube video searches) to help students prepare for these topics.
   - Format: "Search YouTube for: 'Data Structures in Python full course'"
   - Format: "Read documentation: 'Python 'list' methods'"

Format your response clearly using Markdown (headers, subheaders, bullet points, and code blocks for code).
Be SPECIFIC, PRACTICAL, and ensure all questions are directly relevant to the announcement content.
"""
        
        try:
            response = self.gemini_model.generate_content(prompt)
            return response.text.strip()
        except Exception as e:
            return f"⚠️ Error generating practice questions: {str(e)}"


# ---------------------- Helper Function (from Monitor) ----------------------

def _ask_to_save_md(content: str, course_name: str, file_type: str, console: Console):
    """Asks the user if they want to save content to a .md file."""
    
    # Sanitize course name for a safe filename
    safe_course_name = re.sub(r'[^\w\-_\. ]', '_', course_name).replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d-%H%M')
    default_filename = f"{file_type}-{safe_course_name}-{timestamp}.md"

    if typer.confirm(f"\n💾 Do you want to save these {file_type.replace('-', ' ')}?"):
        filename = typer.prompt(
            "Enter filename", 
            default=default_filename
        )
        try:
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
            console.print(f"[green]✓ Successfully saved to [bold]{filename}[/bold][/green]")
        except Exception as e:
            console.print(f"[red]⚠️  Error saving file: {e}[/red]")
        print() # Add a newline for spacing


# ---------------------- Typer Commands (Merged) ----------------------

@app.command()
def list_courses(
    credentials: str = typer.Option("credentials.json", help="Path to credentials file."),
    token: str = typer.Option("token.pickle", help="Path to saved token file."),
):
    """List all your active Google Classroom courses."""
    console.print("📚 [bold]Fetching your courses...[/bold]\n")
    cli = ClassroomBuddyCLI(credentials_file=credentials, token_file=token)
    courses = cli.get_courses()
    if not courses:
        console.print("No active courses found.")
        raise typer.Exit()
    
    console.print(Panel(
        "\n".join([f"• [green]{c['name']}[/green] (ID: {c['id']})" for c in courses]),
        title="Active Courses",
        border_style="blue"
    ))

# --- Command from Study Buddy (Renamed) ---
@app.command("detect-materials")
def detect_materials(
    course_id: str = typer.Option(None, "--course-id", help="Specific course ID to scan."),
    all_courses: bool = typer.Option(False, "--all-courses", help="Scan all available courses."),
    since: int = typer.Option(24, help="Only scan materials from last N hours."),
    credentials: str = typer.Option("credentials.json", help="Path to credentials file"),
    token: str = typer.Option("token.pickle", help="Path to saved token file.")
):
    """
    Detect new lecture materials (PDFs, Docs) and generate study aids.
    """
    console.print("🚀 [bold]Initializing Study Buddy (Material Detector)...[/bold]")
    cli = ClassroomBuddyCLI(credentials_file=credentials, token_file=token)
    print()

    # Determine which courses to process
    if course_id:
        try:
            course_info = cli.classroom_service.courses().get(id=course_id).execute()
            course_name = course_info.get('name', f'Course {course_id}')
            courses = [{'id': course_id, 'name': course_name}]
        except HttpError:
            console.print(f"[red]Error: Could not find course with ID {course_id}[/red]")
            raise typer.Exit()
    elif all_courses:
        courses = cli.get_courses()
        if not courses:
            console.print("[red]No active courses found.[/red]")
            raise typer.Exit()
    else:
        console.print("[yellow]❗ Please specify either --course-id or --all-courses.[/yellow]")
        raise typer.Exit()

    # --- Main Processing Loop ---
    total_materials_found = 0
    for course in courses:
        course_name = course['name']
        console.rule(f"[bold blue]🔍 Scanning for Materials: {course_name}[/bold blue]")
        
        new_materials = cli.get_new_materials(course['id'], since_hours=since)
        
        if not new_materials:
            console.print(f"[dim]No new materials found in the last {since} hours.[/dim]\n")
            continue
            
        console.print(f"📢 Found [green]{len(new_materials)}[/green] new material(s)!")
        total_materials_found += len(new_materials)
        
        for material in new_materials:
            title = material.get('title', 'Untitled Material')
            console.print(f"\n[bold]Processing Material:[/bold] [cyan]{title}[/cyan]")
            
            if not material.get('materials'):
                console.print("[dim]  > No files attached. Skipping.[/dim]")
                continue

            # --- Text Extraction ---
            full_lecture_text = ""
            for item in material.get('materials', []):
                drive_file = item.get('driveFile', {}).get('driveFile')
                if drive_file:
                    text = cli.get_drive_file_text(drive_file)
                    if text:
                        full_lecture_text += f"\n\n--- (Source: {drive_file.get('title', 'N/A')}) ---\n{text}"
            
            if not full_lecture_text.strip():
                console.print("[yellow]  > Could not extract any text from attachments. Skipping.[/yellow]")
                continue
            
            console.print(f"[green]  ✓ Extracted {len(full_lecture_text.split())} words of text.[/green]")
            
            # --- Interactive Generation ---
            if not typer.confirm(f"\n🧠 Do you want to generate study aids for '{title}'?"):
                console.print("[dim]  > Skipping generation.[/dim]")
                continue

            base_filename = f"{re.sub(r'[^\w\-_\. ]', '', title).replace(' ', '_')}"
            
            # Audio Summary (MP3)
            if typer.confirm("  1. Generate an Audio Summary (MP3)?"):
                with Status("[bold magenta]Generating audio summary...[/bold magenta]", console=console):
                    try:
                        # 1. Generate text narration
                        narration_text = cli.generate_audio_narration(full_lecture_text)
                        
                        # 2. Generate MP3 in memory
                        audio_fp = io.BytesIO()
                        tts = gTTS(text=narration_text, lang='en')
                        tts.write_to_fp(audio_fp)
                        
                        # 3. Upload the in-memory MP3 file
                        filename = f"Summary-{base_filename}.mp3"
                        link = cli._upload_audio_to_drive(audio_fp, filename)
                        
                        console.print(f"  [green]✓ Audio Summary generated![/green] [dim]({link})[/dim]")
                    except Exception as e:
                        console.print(f"  [red]Error generating audio: {e}[/red]")
                    
            # Flashcards
            if typer.confirm("  2. Generate Flashcards?"):
                with Status("[bold magenta]Generating flashcards...[/bold magenta]", console=console):
                    content = cli.generate_flashcards(full_lecture_text)
                    filename = f"Flashcards-{base_filename}.csv"
                    link = cli._upload_to_drive(content, filename, mime_type='text/csv')
                console.print(f"  [green]✓ Flashcards generated![/green] [dim]({link})[/dim]")

            # Quiz
            if typer.confirm("  3. Generate a Quiz?"):
                with Status("[bold magenta]Generating quiz...[/bold magenta]", console=console):
                    content = cli.generate_quiz(full_lecture_text)
                    filename = f"Quiz-{base_filename}.md"
                    link = cli._upload_to_drive(content, filename)
                console.print(f"  [green]✓ Quiz generated![/green] [dim]({link})[/dim]")

            console.print(f"\n[bold green]✅ Done processing '{title}'![/bold green]")
            
    console.rule("[bold]Material Detection Complete[/bold]")
    console.print(f"✅ Scanned [bold]{len(courses)}[/bold] course(s).")
    console.print(f"🎯 Found [bold]{total_materials_found}[/bold] new materials.")
    console.print()


# --- Command from Monitor Script (Renamed) ---
@app.command("summarize-announcements")
def summarize_announcements(
    course_id: str = typer.Option(None, "--course-id", help="Specific course ID to summarize."),
    all_courses: bool = typer.Option(False, "--all-courses", help="Summarize all available courses."),
    max_announcements: int = typer.Option(10, "--max", help="Max announcements per course."),
    since: int = typer.Option(None, help="Only include announcements from last N hours."),
    no_summary: bool = typer.Option(False, help="Disable AI summarization, just list."),
    credentials: str = typer.Option("credentials.json", help="Path to credentials file"),
    token: str = typer.Option("token.pickle", help="Path to saved token file.")
):
    """Fetch and summarize all announcements for each course."""
    console.print("🚀 [bold]Initializing Google Classroom Summarizer...[/bold]")
    cli = ClassroomBuddyCLI(credentials_file=credentials, token_file=token)
    print()

    # Determine which courses to process
    if course_id:
        try:
            course_info = cli.classroom_service.courses().get(id=course_id).execute()
            course_name = course_info.get('name', f'Course {course_id}')
        except HttpError:
            course_name = f'Course {course_id}'
        courses = [{'id': course_id, 'name': course_name}]
    elif all_courses:
        courses = cli.get_courses()
        if not courses:
            console.print("[red]No active courses found.[/red]")
            raise typer.Exit()
    else:
        console.print("[yellow]❗ Please specify either --course-id or --all-courses.[/yellow]")
        raise typer.Exit()

    total = 0
    for course in courses:
        console.rule(f"[bold blue]🔍 Summarizing: {course['name']} (ID: {course['id']})[/bold blue]")
        anns = cli.get_announcements(course['id'], max_results=max_announcements, since_hours=since)
        
        if not anns:
            msg = "No announcements d"
            if since:
                msg += f" in the last {since} hours"
            console.print(f"[dim]{msg} in this course.[/dim]\n")
            continue

        total += len(anns)
        console.print(f"📢 Found [green]{len(anns)}[/green] announcement(s):\n")
        for i, ann in enumerate(anns, start=1):
            timestamp = cli._parse_timestamp(ann.get('updateTime')).strftime("%Y-%m-%d %H:%M")
            text = ann.get('text', 'No content').strip().split('\n')[0] # Show first line
            console.print(f"{i}. [{timestamp}] {text[:100]}...") # Truncate long lines

        if not no_summary:
            console.print("\n🤖 [cyan]Generating overall course summary...[/cyan]\n")
            summary = cli.summarize_course_announcements(course['name'], anns)
            console.print(Panel(
                summary,
                title=f"[bold white]📘 COURSE SUMMARY: {course['name']}[/bold white]",
                border_style="blue",
                padding=(1, 2)
            ))
        # Ask to save the summary
            _ask_to_save_md(summary, course['name'], "summary", console)
            
        print()

    console.rule("[bold]Summary Complete[/bold]")
    console.print(f"✅ Processed [bold]{total}[/bold] announcements across [bold]{len(courses)}[/bold] course(s).\n")


# --- Command from Monitor Script (Renamed) ---
@app.command("detect-announcements")
def detect_announcements(
    course_id: str = typer.Option(None, "--course-id", help="Specific course ID to scan."),
    allourses: bool = typer.Option(False, "--all-courses", help="Scan all available courses."),
    max_announcements: int = typer.Option(20, "--max", help="Max announcements to scan per course."),
    since: int = typer.Option(None, help="Only scan announcements from last N hours."),
    keywords_only: bool = typer.Option(False, help="Only show detected keywords without generating ideas."),
    credentials: str = typer.Option("credentials.json", help="Path to credentials file"),
    token: str = typer.Option("token.pickle", help="Path to saved token file.")
):
    """
    Detect project/lab test announcements and generate tailored ideas/questions.
    """
    console.print("🚀 [bold]Initializing Google Classroom Project Detector...[/bold]")
    cli = ClassroomBuddyCLI(credentials_file=credentials, token_file=token)
    print()

    # Determine which courses to process
    if course_id:
        try:
            course_info = cli.classroom_service.courses().get(id=course_id).execute()
            course_name = cour_info.get('name', f'Course {course_id}')
        except HttpError:
            course_name = f'Course {course_id}'
        courses = [{'id': course_id, 'name': course_name}]
    elif all_courses:
        courses = cli.get_courses()
        if not courses:
            console.print("[red]No active courses found.[/red]")
            raise typer.Exit()
    else:
        console.print("[yellow]❗ Please specify either --course-id or --all-courses.[/yellow]")
        raise typer.Exit()

    total_announcements = 0
    total_projects_detected = 0
    total_lab_tests_detected = 0
    
    for course in courses:
        console.rule(f"[bold blue]🔍 Scanning for Announcements: {course['name']}[/bold blue]")
        anns = cli.get_announcements(course['id'], max_results=max_announcements, since_hours=since)
        
        if not anns:
            msg = "No announcements found"
            if since:
                msg += f" in the last {since} hours"
            console.print(f"[dim]{msg} in this course.[/dim]\n")
            continue

        total_announcements += len(anns)
        console.print(f"📢 Scanning [green]{len(anns)}[/green] announcement(s) for keywords...\n")
        
        # Detect project announcements
        project_anns = cli.detect_project_announcements(anns)
        
        # Detect lab test announcements
        lab_test_anns = cli.detect_lab_test_announcements(anns)
        
        if not project_anns and not lab_test_anns:
            console.print("[dim]✗ No project or lab test announcem detected in this course.[/dim]\n")
            continue
        
        # --- Process Project Announcements ---
        if project_anns:
            total_projects_detected += len(project_anns)
            console.print(f"🎯 [bold green]Found {len(project_anns)} project-related announcement(s)![/bold green]\n")
            
            for idx, (ann, keywords) in enumerate(project_anns, start=1):
                timestamp = cli._parse_timestamp(ann.get('updateTime')).strftime("%Y-%m-%d %H:%M")
                text = ann.get('text', 'No content').strip()
                
                console.print(Panel(
                    f"[bold yellow]📅 Posted:[/bold yellow] {timestamp}\n\n"
                    f"[bold yellow]🔑 Detected Keywords:[/bold yellow] {', '.join(keywords)}\n\n"
                    f"[bold yellow]📝 Full Announcement:[/bold yellow]\n{text}",
                    title=f"[bold cyan]Project Announcement #{idx} - {course['name']}[/bold cyan]",
                    border_style="cyan",
              padding=(1, 2)
                ))
                
                if keywords_only:
                    console.print()
                    continue
                
                console.print("\n💡 [bold magenta]Analyzing announcement and generating tailored project ideas...[/bold magenta]\n")
                
                project_ideas = cli.generate_tailored_project_ideas(
                    course['name'], 
                    text, 
                    keywords
                )
              
                console.print(Panel(
                    Markdown(project_ideas),
                    title=f"[bold green]🚀 Tailored Project Ideas & Resources[/bold green]",
                    border_style="green",
                    padding=(1, 2)
                ))
                
                # Ask to save the project ideas
                _ask_to_save_md(project_ideas, course['name'], f"project-ideas-{idx}", console)
                
                console.print("\n" + "="*80 + "\n")
        # --- Process Lab Test Announcements ---
        if lab_test_anns:
            total_lab_tests_detected += len(lab_test_anns)
            console.print(f"🧪 [bold yellow]Found {len(lab_test_anns)} lab test/evaluation announcement(s)![/bold yellow]\n")
            
            for idx, (ann, keywords) in enumerate(lab_test_anns, start=1):
                timestamp = cli._parse_timestamp(ann.get('updateTime')).strftime("%Y-%m-%d %H:%M")
                text = ann.get('text', 'No content').strip()
              
                console.print(Panel(
                    f"[bold yellow]📅 Posted:[/bold yellow] {timestamp}\n\n"
                    f"[bold yellow]🔑 Detected Keywords:[/bold yellow] {', '.join(keywords)}\n\n"
                    f"[bold yellow]📝 Full Announcement:[/bold yellow]\n{text}",
                    title=f"[bold magenta]Lab Test Announcement #{idx} - {course['name']}[/bold magenta]",
                    border_style="magenta",
                    padding=(1, 2)))
                             
                if keywords_only:
                    console.print()
                    continue
                    
                console.print("\n💡 [bold cyan]Analyzing announcement and generating practice questions...[/bold cyan]\n")
                
                practice_questions = cli.generate_practice_questions(
                    course['name'], 
                    text, 
                    keywords
                )
                
                console.print(Panel(Markdown(practice_questions),
                    title=f"[bold blue]📚 Practice Questions & Study Guide[/bold blue]",
                    border_style="blue",
                    padding=(1, 2)
                ))
                
                # Ask to save the practice questions
                _ask_to_save_md(practice_questions, course['name'], f"practice-questions-{idx}", console)
                
                console.print("\n" + "="*80 + "\n")

    # Summary
    console.rule("[bold]Announcement Detection Complete[/bold]")
    console.print(f"✅ Scanned [bold]{total_announcements}[/bold] announcements across [bold]{len(courses)}[/bold] course(s)")
    console.print(f"🎯 Detected [bold green]{total_projects_detected}[/bold green] project-related announcements")
    console.print(f"🧪 Detected [bold yellow]{total_lab_tests_detected}[/bold yellow] lab test/evaluation announcements")
    
    if total_projects_detected == 0 and total_lab_tests_detected == 0:
        console.print(f"[yellow]💡 Tip: Try increasing the scan window with --since option or check more courses with --all-courses[/yellow]")
    
    console.print()


# --- Command from Monitor Script (Renamed) ---
@app.command("analyze-announcement")
def analyze_announcement(
    announcement_text: str = typer.Argument(..., help="Project announcement text to analyze"),
    course_name: str = typer.Option("General", help="Course name for context"),
    credentials: str = typer.Option("credentials.json", help="Path to credentia file."),
    token: str = typer.Option("token.pickle", help="Path to saved token file.")
):
    """
    Analyze specific announcement text and generate tailored project/lab ideas.
    """
    console.print(f"🔍 [bold]Analyzing announcement text for: {course_name}[/bold]\n")
    # Authentication is needed just to set up the Gemini model
    cli = ClassroomBuddyCLI(credentials_file=credentials, token_file=token)
    
    # Detect project keywords
    project_keywords = []
    text_lower = announcement_textower()
    for keyword in PROJECT_KEYWORDS:
        if re.search(r'\b' + re.escape(keyword) + r'\b', text_lower, re.IGNORECASE):
            project_keywords.append(keyword)
    
    # Detect lab test keywords
    lab_test_keywords = []
    for keyword in LAB_TEST_KEYWORDS:
        if re.search(r'\b' + re.escape(keyword) + r'\b', text_lower, re.IGNORECASE):
            lab_test_keywords.append(keyword)

    if not project_keywords and not lab_test_keywords:
        console.print("[yellow]⚠️ No project ob test keywords detected in the provided text.[/yellow]")
        console.print(f"[dim]Project keywords: {', '.join(PROJECT_KEYWORDS[:5])}...[/dim]")
        console.print(f"[dim]Lab Test keywords: {', '.join(LAB_TEST_KEYWORDS[:5])}...[/dim]\n")
        return

    # Prioritize Projects: If project keywords are present, run project analysis
    if project_keywords:
        console.print(Panel(
            f"[bold yellow]🔑 Detected Project Keywords:[/bold yellow] {', '.join(project_keywords)}\n\n"
            f"[bold yellow]📝 Announcement Text:[/bold yellow]\n{announcement_text}",
            title=f"[bold cyan]Project Announcement Analysis[/bold cyan]",
            border_style="cyan",
            padding=(1, 2)
        ))
        
        console.print("\n💡 [bold magenta]Generating tailored project ideas...[/bold magenta]\n")
        
        project_ideas = cli.generate_tailored_project_ideas(
            course_name, 
            announcement_text, 
            project_keywords
        )
        
        console.print(Panel(
            Markdown(project_ideas),
            title=f"[bold green]🚀 Tailored Project Ideas & Resources[/bold green]",
            border_style="green",
            padding=(1, 2)
        ))
        
        # Ask to save the project ideas
        _ask_to_save_md(project_ideas, course_name, "project-ideas-manual", console)
    
    # Else If: If no project keywords, but lab test keywords are present, run lab test analysis
    elif lab_test_keywords:
        console.print(Panel(
            f"[bold yellow]🔑 Detected Lab Test Keywords:[/bold yellow] {', '.join(lab_test_keywords)}\n\n"
            f"[bold yellow]📝 Announcement Text:[/bold yellow]\n{announcement_text}",
            title=f"[bold magenta]Lab Test Announcement Analysis[/bold magenta]",
            border_style="magenta",
            padding=(1, 2)
        ))
        
        console.print("\n💡 [bold cyan]Analyzing announcement and generating practice questions...[/bold cyan]\n")
        
        practice_questions = cli.generate_practice_questions(
            course_name, 
            announcement_text, 
            lab_test_keywords
        )
        
        console.print(Panel(
            Markdown(practice_questions),
            title=f"[bold blue]📚 Practice Questions & Study Guide[/bold blue]",
            border_style="blue",
            padding=(1, 2)
        ))
        
        # Ask to save the practice questions
        _ask_to_save_md(practice_questions, course_name, "practice-questions-manual", console)
        
    console.print()

# --- NEW FEATURE: Generate .docx from code ---
@app.command("generate-doc")
def generate_doc(
    root_dir: str = typer.Option("src", "--source", help="Source code directory to scan."),
    output_dir: str = typer.Option("output", "--output-dir", help="Directory to save the generated document."),
    output_file_name: str = typer.Option("Code_Documentation", "--filename", help="Name of the output file (without .docx extension)."),
    heading: str = typer.Option("Project Code Documentation", "--title", help="Main heading of the document."),
    code_font: str = typer.Option("Courier New", "--font", help="Font for code blocks."),
    extensions_str: str = typer.Option(".py,.js,.html,.css,.md,.java,.c,.cpp", "--extensions", help="Comma-separated list of file extensions to include."),
    project_name: str = typer.Option("My Project", "--project-name", help="Project name for the footer."),
    github_link: str = typer.Option("https://github.com/user/repo", "--github", help="GitHub repository link for the footer."),
):
    """
    Scans a source code directory and generates a formatted .docx file.
    """
    console.print("📄 [bold]Initializing Code-to-Docx Generator...[/bold]")

    if not DOCX_AVAILABLE:
        console.print("[red]Error: `python-docx` is not installed.[/red]")
        console.print("[yellow]Please install it by running: pip install python-docx[/yellow]")
        raise typer.Exit()
        
    if not os.path.isdir(root_dir):
        console.print(f"[red]Error: Souce directory '{root_dir}' not found.[/red]")
        raise typer.Exit()

    try:
        os.makedirs(output_dir, exist_ok=True)
    except OSError as e:
        console.print(f"[red]Error creating output directory '{output_dir}': {e}[/red]")
        raise typer.Exit()
        
    extensions = [ext.strip() for ext in extensions_str.split(',')]
    output_file = os.path.join(output_dir, f"{output_file_name}.docx")
    
    doc = Document()
    
    # Add and format the main title
    title = doc.add_heading(heading, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    console.print(f"🔍 Scanning [cyan]'{root_dir}'[/cyan] for files with extensions: [yellow]{', '.join(extensions)}[/yellow]")
    
    files_processed = 0
    # Traverse the file structure
    for folder, _, files in os.walk(root_dir):
        for filename in files:
            if any(filename.endswith(ext) for ext in extensions):
                files_processed += 1
                file_path = os.path.join(folder, filename)
                rel_path = os.path.relpath(file_path, root_dir)

                console.print(f"  > Adding file: [dim]{rel_path}[/dim]")

                # Add Section Header for the file
                doc.add_heading(f"File: {rel_path}", level=2)

                # Add the code content to the document
                try:
                    with open(file_path, "r", encoding='utf-8', errors='ignore') as f:
                        content = f.read()

                    para = doc.add_paragraph()
                    run = para.add_run(content)
                    run.font.name = code_font
                    run.font.size = Pt(10)
                except Exception as e:
                    console.print(f"    [red]Could not read file '{rel_path}': {e}[/red]")
    
    if files_processed == 0:
        console.print("[yellow]Warning: No matching files were found to add to the document.[/yellow]")
    
    # Add a new section for the footer to apply it only to the last page
    doc.add_section()
    
    # Add footer
    end_footer = doc.sections[-1].footer
    end_footer_para = end_footer.paragraphs[0]
    end_footer_para.text = f"© {datetime.now().year} – {project_name} – {github_link}"
    end_footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    end_footer_para.runs[0].font.size = Pt(10)

    # Save the document
    try:
        doc.save(output_file)
        console.print(f"\n[bold green]✅ Document successfully generated![/bold green]")
        console.print(f"📄 Saved to: [cyan]{output_file}[/cyan]")
    except Exception as e:
        console.print(f"[red]Error saving document: {e}[/red]")
    console.print()

# ---------------------- Main ----------------------
if __name__ == "__main__":
    app()
